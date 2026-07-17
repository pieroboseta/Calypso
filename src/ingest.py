from pathlib import Path
import json
import hashlib
import pickle
import sys

from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader


if getattr(sys, "frozen", False):
    PROJECT_FOLDER = Path(sys.executable).resolve().parent
else:
    SRC_FOLDER = Path(__file__).resolve().parent
    PROJECT_FOLDER = SRC_FOLDER.parent

LIBRARY_FOLDER = PROJECT_FOLDER / "library"
DATA_FOLDER = PROJECT_FOLDER / "data"
DOCUMENTS_FILE = DATA_FOLDER / "documents.pkl"
CACHE_FOLDER = DATA_FOLDER / "document_cache"
MANIFEST_FILE = CACHE_FOLDER / "manifest.json"

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def _file_signature(
    path: Path,
    previous_signature: dict | None = None,
) -> dict:
    stat = path.stat()
    basic = {
        "size": stat.st_size,
        "modified": stat.st_mtime_ns,
    }

    if (
        previous_signature
        and previous_signature.get("size") == basic["size"]
        and previous_signature.get("modified") == basic["modified"]
        and previous_signature.get("sha256")
    ):
        basic["sha256"] = previous_signature["sha256"]
        return basic

    digest = hashlib.sha256()
    with path.open("rb") as file:
        for block in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(block)

    basic["sha256"] = digest.hexdigest()
    return basic


def _cache_file_name(source_path: Path) -> str:
    safe_name = "".join(
        character if character.isalnum() else "_"
        for character in source_path.name
    )
    return f"{safe_name}.pkl"


def _load_manifest() -> dict:
    if not MANIFEST_FILE.exists():
        return {}

    try:
        data = json.loads(MANIFEST_FILE.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError):
        return {}


def _save_manifest(manifest: dict) -> None:
    MANIFEST_FILE.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def _read_pdf(
    pdf_path: Path,
    progress_callback=None,
) -> list[Document]:
    reader = PdfReader(str(pdf_path))
    documents = []
    total_pages = len(reader.pages)

    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text()

        if text and text.strip():
            documents.append(
                Document(
                    page_content=text.strip(),
                    metadata={
                        "book": pdf_path.name,
                        "page": page_number,
                        "file_type": "pdf",
                    },
                )
            )

        if progress_callback:
            percent = int(page_number / max(total_pages, 1) * 100)
            progress_callback(
                percent,
                f"Reading {pdf_path.name}: "
                f"{page_number}/{total_pages} pages",
            )

    return documents


def _read_docx(
    docx_path: Path,
    progress_callback=None,
) -> list[Document]:
    document = DocxDocument(str(docx_path))
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    table_texts = []
    for table in document.tables:
        for row in table.rows:
            values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]
            if values:
                table_texts.append(" | ".join(values))

    all_blocks = paragraphs + table_texts

    if not all_blocks:
        return []

    documents = []
    block_size = 40
    total_blocks = len(all_blocks)

    for start in range(0, total_blocks, block_size):
        end = min(start + block_size, total_blocks)
        text = "\n\n".join(all_blocks[start:end]).strip()

        if text:
            section_number = start // block_size + 1
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "book": docx_path.name,
                        "page": section_number,
                        "file_type": "docx",
                    },
                )
            )

        if progress_callback:
            percent = int(end / total_blocks * 100)
            progress_callback(
                percent,
                f"Reading {docx_path.name}: "
                f"{end}/{total_blocks} blocks",
            )

    return documents


def _read_txt(
    txt_path: Path,
    progress_callback=None,
) -> list[Document]:
    text = None

    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            text = txt_path.read_text(encoding=encoding)
            break
        except UnicodeDecodeError:
            continue

    if text is None or not text.strip():
        return []

    lines = [line.rstrip() for line in text.splitlines()]
    documents = []
    line_group_size = 120
    total_lines = max(len(lines), 1)

    for start in range(0, len(lines), line_group_size):
        end = min(start + line_group_size, len(lines))
        block = "\n".join(lines[start:end]).strip()

        if block:
            section_number = start // line_group_size + 1
            documents.append(
                Document(
                    page_content=block,
                    metadata={
                        "book": txt_path.name,
                        "page": section_number,
                        "file_type": "txt",
                    },
                )
            )

        if progress_callback:
            percent = int(end / total_lines * 100)
            progress_callback(
                percent,
                f"Reading {txt_path.name}: "
                f"{end}/{len(lines)} lines",
            )

    return documents


def _read_file(
    source_path: Path,
    progress_callback=None,
) -> list[Document]:
    suffix = source_path.suffix.lower()

    if suffix == ".pdf":
        return _read_pdf(source_path, progress_callback)

    if suffix == ".docx":
        return _read_docx(source_path, progress_callback)

    if suffix == ".txt":
        return _read_txt(source_path, progress_callback)

    return []


def ingest_library(progress_callback=None):
    DATA_FOLDER.mkdir(parents=True, exist_ok=True)
    LIBRARY_FOLDER.mkdir(parents=True, exist_ok=True)
    CACHE_FOLDER.mkdir(parents=True, exist_ok=True)

    library_files = sorted(
        path
        for path in LIBRARY_FOLDER.iterdir()
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )

    if not library_files:
        raise RuntimeError(
            "No PDF, DOCX, or TXT files were found in the library folder."
        )

    old_manifest = _load_manifest()
    new_manifest = {}
    all_documents = []
    total_files = len(library_files)

    for file_index, source_path in enumerate(library_files, start=1):
        cached_entry = old_manifest.get(source_path.name, {})
        signature = _file_signature(
            source_path,
            cached_entry.get("signature"),
        )
        cache_name = _cache_file_name(source_path)
        cache_path = CACHE_FOLDER / cache_name

        cache_is_valid = (
            cached_entry.get("signature") == signature
            and cached_entry.get("cache_file") == cache_name
            and cache_path.exists()
        )

        file_documents = []

        if cache_is_valid:
            try:
                with cache_path.open("rb") as file:
                    file_documents = pickle.load(file)

                if progress_callback:
                    overall = int(file_index / total_files * 100)
                    progress_callback(
                        overall,
                        f"Using cached text: {source_path.name}",
                    )
            except (OSError, pickle.PickleError, EOFError):
                cache_is_valid = False

        if not cache_is_valid:
            def file_progress(percent, message):
                if progress_callback:
                    completed = file_index - 1
                    overall = int(
                        ((completed + percent / 100) / total_files) * 100
                    )
                    progress_callback(overall, message)

            file_documents = _read_file(
                source_path,
                progress_callback=file_progress,
            )

            with cache_path.open("wb") as file:
                pickle.dump(file_documents, file)

        all_documents.extend(file_documents)

        new_manifest[source_path.name] = {
            "signature": signature,
            "cache_file": cache_name,
            "sections": len(file_documents),
            "file_type": source_path.suffix.lower().lstrip("."),
        }

    valid_cache_files = {
        entry["cache_file"]
        for entry in new_manifest.values()
        if entry.get("cache_file")
    }

    for cache_path in CACHE_FOLDER.glob("*.pkl"):
        if cache_path.name not in valid_cache_files:
            try:
                cache_path.unlink()
            except OSError:
                pass

    if not all_documents:
        raise RuntimeError(
            "The library files were found, but no readable text was extracted."
        )

    with DOCUMENTS_FILE.open("wb") as file:
        pickle.dump(all_documents, file)

    _save_manifest(new_manifest)

    if progress_callback:
        progress_callback(
            100,
            f"Loaded {len(all_documents)} searchable sections",
        )

    return all_documents


if __name__ == "__main__":
    ingest_library(
        progress_callback=lambda percent, message: print(
            f"{percent}% - {message}"
        )
    )